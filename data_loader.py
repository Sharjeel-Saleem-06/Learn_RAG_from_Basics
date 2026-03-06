"""
============================================
FILE 1 OF 8: DATA_LOADER.PY — Document Loading
============================================

📝 PURPOSE:
    This is the FIRST step in the RAG pipeline.
    Before AI can answer questions about your documents, it needs to READ them.
    
    This file handles loading documents from various formats:
    - PDF files (.pdf)
    - Text files (.txt) 
    - Word documents (.docx)
    
    Think of it as the "eyes" of your RAG system — it reads the raw documents.

🎓 CONCEPTS COVERED:
    1. Document Parsing — Extracting text from different file formats
    2. Document Representation — How to structure loaded documents as Python objects
    3. Metadata — Extra information ABOUT the document (filename, page number, etc.)
    4. Error Handling — Gracefully handling corrupted or unsupported files
    5. Factory Pattern — Choosing the right loader based on file type

🔗 HOW IT CONNECTS:
    data_loader.py  →  chunker.py  →  embedding_generator.py  →  vector_db.py
    ^^^^^^^^^^^^^^
    YOU ARE HERE
    
    Output of this file → Input of chunker.py
    We load raw text here, then chunker.py splits it into smaller pieces.

🏗️ ARCHITECTURE DECISION:
    We use a "Document" dataclass to wrap text + metadata.
    Why not just pass raw strings around?
    
    Because in a REAL system, you need to know:
    - WHERE did this text come from? (which file?)
    - WHICH PAGE was it on? (for citation)
    - WHEN was it loaded? (for freshness)
    
    This metadata is CRITICAL for production RAG systems.
    It lets you cite sources: "According to page 5 of report.pdf..."

💡 INTERVIEW QUESTION: "How do you handle different document types in RAG?"
    ANSWER: Use a Document abstraction that normalizes all formats into
    a common structure (text + metadata). Implement format-specific loaders
    that all produce the same output type. This follows the "Strategy Pattern"
    — different algorithms (PDF parsing, text reading) behind a unified interface.

============================================
"""

import os
from dataclasses import dataclass, field
from typing import List, Optional
from datetime import datetime


# ============================================
# DOCUMENT DATACLASS
# ============================================
# This is the "universal format" for all documents in our system.
# Whether the original was PDF, txt, or docx, after loading
# it becomes a Document object.
#
# 💡 WHY A DATACLASS?
# Instead of passing around raw strings and losing track of
# where text came from, we wrap everything in a clean object:
#
#   doc = Document(
#       text="Machine learning is...",
#       metadata={"source": "ml_book.pdf", "page": 3}
#   )
#
# Now anywhere in your code, you can access:
#   doc.text     → The actual content
#   doc.metadata → Where it came from
# ============================================

@dataclass
class Document:
    """
    Represents a loaded document with its text content and metadata.
    
    Attributes:
        text (str): The raw text content extracted from the document.
        metadata (dict): Information about the document:
            - source: filename or path
            - page: page number (for PDFs)
            - file_type: pdf, txt, docx
            - loaded_at: timestamp when document was loaded
            - total_pages: total pages in the source document
    
    💡 REAL-WORLD NOTE:
        In production systems (like LangChain), the Document class is
        almost identical to this. LangChain's Document has:
        - page_content (str): same as our 'text'
        - metadata (dict): same as our 'metadata'
        
        So you're learning the REAL pattern used in industry!
    """
    text: str
    metadata: dict = field(default_factory=dict)
    
    def __len__(self):
        """Returns the length of the text content."""
        return len(self.text)
    
    def __repr__(self):
        """Nice string representation for debugging."""
        preview = self.text[:100] + "..." if len(self.text) > 100 else self.text
        source = self.metadata.get("source", "unknown")
        return f"Document(source='{source}', length={len(self.text)}, preview='{preview}')"


# ============================================
# PDF LOADER
# ============================================
# PDFs are the most common document format in enterprise RAG systems.
# 
# 💡 WHY IS PDF PARSING HARD?
# PDFs were designed for PRINTING, not for text extraction.
# A PDF is basically a set of instructions like:
#   "Draw the letter 'H' at position (72, 300)"
#   "Draw the letter 'e' at position (80, 300)"
# 
# So extracting "Hello" requires figuring out that these letters
# are close together and form a word. This is why PDF parsers
# sometimes mess up spacing, tables, and multi-column layouts.
#
# 💡 INTERVIEW INSIGHT:
# In production, people often use specialized PDF parsers:
# - PyPDF2: Good for simple PDFs (what we're using)
# - pdfplumber: Better for PDFs with tables
# - Unstructured: Best for complex layouts (handles images, tables)
# - Amazon Textract / Google Document AI: Cloud-based, most accurate
# ============================================

def load_pdf(file_path: str) -> List[Document]:
    """
    Load a PDF file and return a list of Document objects (one per page).
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        List[Document]: One Document per page, with page-specific metadata
        
    Why one Document per page?
        - Enables page-level citations ("See page 5")
        - Different pages may cover different topics
        - Keeps individual documents reasonably sized
        
    Example:
        >>> docs = load_pdf("machine_learning_book.pdf")
        >>> print(docs[0].text[:200])  # First 200 chars of page 1
        >>> print(docs[0].metadata)    # {'source': 'machine_learning_book.pdf', 'page': 1, ...}
    """
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError(
            "PyPDF2 is not installed! Run: pip install PyPDF2\n"
            "PyPDF2 is the library that reads PDF files."
        )
    
    # Validate the file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"❌ File not found: {file_path}")
    
    if not file_path.lower().endswith('.pdf'):
        raise ValueError(f"❌ Not a PDF file: {file_path}")
    
    documents = []
    filename = os.path.basename(file_path)
    
    # ============================================
    # PdfReader reads the PDF structure
    # Think of it as opening the PDF and getting
    # access to each page
    # ============================================
    reader = PdfReader(file_path)
    total_pages = len(reader.pages)
    
    print(f"📄 Loading PDF: {filename} ({total_pages} pages)")
    
    for page_num, page in enumerate(reader.pages, start=1):
        # ============================================
        # extract_text() pulls the text content from a page
        # 
        # ⚠️ LIMITATION: This won't extract:
        # - Text inside images (need OCR for that)
        # - Complex table structures (cells may merge)
        # - Headers/footers may be included
        #
        # For a beginner project, this is perfectly fine.
        # Advanced: Use pdfplumber or Unstructured for better extraction.
        # ============================================
        text = page.extract_text()
        
        if text and text.strip():  # Only add non-empty pages
            doc = Document(
                text=text.strip(),
                metadata={
                    "source": filename,
                    "page": page_num,
                    "total_pages": total_pages,
                    "file_type": "pdf",
                    "loaded_at": datetime.now().isoformat(),
                    "file_path": file_path
                }
            )
            documents.append(doc)
        else:
            print(f"  ⚠️ Page {page_num} is empty or contains only images (skipped)")
    
    print(f"  ✅ Loaded {len(documents)} pages with text content")
    return documents


# ============================================
# TEXT FILE LOADER
# ============================================
# The simplest loader — just reads a plain text file.
#
# 💡 ENCODING NOTE:
# We use 'utf-8' encoding which handles most languages.
# If you get encoding errors, try 'latin-1' or 'cp1252'.
#
# What is encoding?
# Computers store text as numbers. 'A' = 65, 'B' = 66, etc.
# UTF-8 is the standard way to map ALL characters (including 
# Arabic, Chinese, emoji 🎉) to numbers.
# ============================================

def load_text(file_path: str) -> List[Document]:
    """
    Load a plain text file and return it as a single Document.
    
    Args:
        file_path (str): Path to the text file
        
    Returns:
        List[Document]: A list containing one Document
        
    Why a list with one item?
        To keep a CONSISTENT interface. Whether you load a PDF (multiple pages)
        or a text file (single document), you always get List[Document] back.
        This makes downstream code simpler — it doesn't need to check types.
        
        💡 INTERVIEW CONCEPT: "Uniform Interface"
        All loaders return the same type, making them interchangeable.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"❌ File not found: {file_path}")
    
    filename = os.path.basename(file_path)
    print(f"📝 Loading text file: {filename}")
    
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    
    if not text.strip():
        print(f"  ⚠️ File is empty!")
        return []
    
    doc = Document(
        text=text.strip(),
        metadata={
            "source": filename,
            "file_type": "txt",
            "loaded_at": datetime.now().isoformat(),
            "file_path": file_path,
            "char_count": len(text)
        }
    )
    
    print(f"  ✅ Loaded {len(text):,} characters")
    return [doc]


# ============================================
# WORD DOCUMENT LOADER
# ============================================
# Loads .docx files (Microsoft Word format)
#
# 💡 FUN FACT: A .docx file is actually a ZIP archive containing:
# - XML files with the text content
# - Images, fonts, styles
# The python-docx library handles unpacking all of this for us.
# ============================================

def load_docx(file_path: str) -> List[Document]:
    """
    Load a Word document (.docx) and return it as a Document.
    
    Args:
        file_path (str): Path to the .docx file
        
    Returns:
        List[Document]: A list containing one Document
    """
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError(
            "python-docx is not installed! Run: pip install python-docx\n"
            "python-docx is the library that reads Word documents."
        )
    
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"❌ File not found: {file_path}")
    
    filename = os.path.basename(file_path)
    print(f"📃 Loading Word document: {filename}")
    
    doc = DocxDocument(file_path)
    
    # ============================================
    # A Word document has "paragraphs" — each is a block of text.
    # We join them with newlines to preserve structure.
    # ============================================
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    text = "\n\n".join(paragraphs)
    
    if not text.strip():
        print(f"  ⚠️ Document is empty!")
        return []
    
    result = Document(
        text=text.strip(),
        metadata={
            "source": filename,
            "file_type": "docx",
            "loaded_at": datetime.now().isoformat(),
            "file_path": file_path,
            "paragraph_count": len(paragraphs)
        }
    )
    
    print(f"  ✅ Loaded {len(paragraphs)} paragraphs, {len(text):,} characters")
    return [result]


# ============================================
# UNIVERSAL LOADER — The "Smart Router"
# ============================================
# This is the main function you'll use. It automatically detects
# the file type and calls the appropriate loader.
#
# 💡 DESIGN PATTERN: "Factory Method"
# Instead of making the user choose which loader to call,
# we have ONE function that figures it out automatically.
#
# User just calls: load_document("any_file.pdf")
# And it works regardless of file type.
# ============================================

# Map file extensions to their loader functions
SUPPORTED_FORMATS = {
    ".pdf": load_pdf,
    ".txt": load_text,
    ".docx": load_docx,
}


def load_document(file_path: str) -> List[Document]:
    """
    🎯 MAIN FUNCTION — Load any supported document type.
    
    Automatically detects the file type and uses the appropriate loader.
    
    Args:
        file_path (str): Path to any supported document
        
    Returns:
        List[Document]: Loaded documents with text and metadata
        
    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file type is not supported
        
    Example:
        >>> docs = load_document("research_paper.pdf")
        >>> for doc in docs:
        ...     print(f"Page {doc.metadata['page']}: {doc.text[:100]}...")
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"❌ File not found: {file_path}")
    
    # Get the file extension (.pdf, .txt, .docx)
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    if ext not in SUPPORTED_FORMATS:
        supported = ", ".join(SUPPORTED_FORMATS.keys())
        raise ValueError(
            f"❌ Unsupported file type: '{ext}'\n"
            f"   Supported formats: {supported}"
        )
    
    # Call the appropriate loader function
    loader_func = SUPPORTED_FORMATS[ext]
    return loader_func(file_path)


def load_directory(dir_path: str) -> List[Document]:
    """
    Load ALL supported documents from a directory.
    
    This is useful when you have a folder of documents you want
    to add to your RAG system all at once.
    
    Args:
        dir_path (str): Path to directory containing documents
        
    Returns:
        List[Document]: All documents from all files in the directory
        
    Example:
        >>> docs = load_directory("./my_documents/")
        >>> print(f"Loaded {len(docs)} documents from directory")
    """
    if not os.path.isdir(dir_path):
        raise NotADirectoryError(f"❌ Not a directory: {dir_path}")
    
    all_documents = []
    supported_extensions = set(SUPPORTED_FORMATS.keys())
    
    print(f"📂 Scanning directory: {dir_path}")
    
    for filename in sorted(os.listdir(dir_path)):
        _, ext = os.path.splitext(filename)
        if ext.lower() in supported_extensions:
            file_path = os.path.join(dir_path, filename)
            try:
                docs = load_document(file_path)
                all_documents.extend(docs)
            except Exception as e:
                print(f"  ❌ Error loading {filename}: {e}")
    
    print(f"\n📊 Total: Loaded {len(all_documents)} documents from {dir_path}")
    return all_documents


# ============================================
# TEST & DEMO SECTION
# ============================================
# Run this file directly to test the data loader:
#   python data_loader.py
#
# It will create a sample text file and demonstrate loading.
# ============================================

if __name__ == "__main__":
    from rich import print as rprint
    from rich.panel import Panel
    from rich.table import Table
    
    rprint(Panel.fit(
        "[bold cyan]📚 DATA LOADER — Test & Demo[/bold cyan]\n\n"
        "This module is responsible for reading documents\n"
        "and converting them into a standard format (Document objects)\n"
        "that the rest of the RAG pipeline can work with.",
        title="File 1 of 8",
        border_style="cyan"
    ))
    
    # Create a sample document directory
    sample_dir = "./sample_docs"
    os.makedirs(sample_dir, exist_ok=True)
    
    # Create a sample text file with educational content
    sample_text = """
Retrieval-Augmented Generation (RAG) is a technique that enhances Large Language Models 
by providing them with relevant context from external documents before generating a response.

The RAG process works in several steps:

1. Document Ingestion: First, you load your documents (PDFs, text files, etc.) into the system.
   This is what the data_loader module handles.

2. Chunking: The loaded text is split into smaller, manageable pieces called "chunks".
   Each chunk should be self-contained enough to be meaningful on its own.

3. Embedding: Each chunk is converted into a numerical vector (embedding) that captures
   the semantic meaning of the text. Similar texts produce similar vectors.

4. Vector Storage: These embeddings are stored in a vector database (like ChromaDB)
   which is optimized for similarity search.

5. Query Processing: When a user asks a question, their query is also converted to an
   embedding vector.

6. Retrieval: The system finds the chunks whose embeddings are most similar to the
   query embedding. These are the "relevant context" documents.

7. Generation: The retrieved chunks are combined with the user's question and sent
   to an LLM, which generates a response grounded in the actual document content.

This approach significantly reduces hallucination because the LLM bases its answers
on real, retrieved information rather than relying solely on its training data.

Key advantages of RAG over fine-tuning:
- No expensive model training required
- Documents can be updated without retraining
- Source citations are possible
- Works with any LLM (OpenAI, Claude, Llama, etc.)
- Cost-effective for most use cases
    """.strip()
    
    sample_file = os.path.join(sample_dir, "rag_introduction.txt")
    with open(sample_file, 'w') as f:
        f.write(sample_text)
    
    rprint(f"\n[yellow]Created sample file: {sample_file}[/yellow]\n")
    
    # ============================================
    # TEST 1: Load a single text file
    # ============================================
    rprint("[bold green]--- Test 1: Load single text file ---[/bold green]")
    docs = load_document(sample_file)
    
    for doc in docs:
        rprint(f"\n[cyan]Document loaded:[/cyan]")
        rprint(f"  Source: {doc.metadata['source']}")
        rprint(f"  Type: {doc.metadata['file_type']}")
        rprint(f"  Characters: {doc.metadata['char_count']:,}")
        rprint(f"  Preview: {doc.text[:200]}...")
    
    # ============================================
    # TEST 2: Load directory
    # ============================================
    rprint(f"\n[bold green]--- Test 2: Load entire directory ---[/bold green]")
    all_docs = load_directory(sample_dir)
    
    # Show results in a nice table
    table = Table(title="Loaded Documents")
    table.add_column("Source", style="cyan")
    table.add_column("Type", style="green")
    table.add_column("Length", style="yellow", justify="right")
    
    for doc in all_docs:
        table.add_row(
            doc.metadata.get("source", "unknown"),
            doc.metadata.get("file_type", "unknown"),
            f"{len(doc.text):,} chars"
        )
    
    rprint(table)
    
    # ============================================
    # TEST 3: Error handling
    # ============================================
    rprint(f"\n[bold green]--- Test 3: Error handling ---[/bold green]")
    try:
        load_document("nonexistent_file.pdf")
    except FileNotFoundError as e:
        rprint(f"[red]Caught expected error:[/red] {e}")
    
    try:
        load_document("some_file.xyz")
    except (ValueError, FileNotFoundError) as e:
        rprint(f"[red]Caught expected error:[/red] {e}")
    
    rprint(Panel.fit(
        "[bold green]✅ Data Loader working correctly![/bold green]\n\n"
        "[yellow]WHAT YOU LEARNED:[/yellow]\n"
        "• Document parsing — reading different file formats\n"
        "• Document representation — wrapping text + metadata\n"
        "• Factory pattern — one function handles all file types\n"
        "• Error handling — graceful failures with clear messages\n\n"
        "[yellow]NEXT STEP:[/yellow] chunker.py — splitting text into chunks",
        title="Summary",
        border_style="green"
    ))
